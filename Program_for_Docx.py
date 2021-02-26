import PySimpleGUI as sg  # интерфейс
import docx  # работа с вордом
import re  # регулярные выражения


# функции, использованные в программе

def numbers(string_with_nums):
    """
    Функция принимает на входе строку
    и выдает список цифр, содержащихся в ней
    """
    return list(map(float, re.findall(r'(?<!\d)-?\d*[.,]?\d+', string_with_nums.replace(',', '.'))))


def more(need_list, have_list):
    """
    Проверяет условие need_num_list[0] <= have_num_list[0]

    Принимает на входе числа, полученные функцией numbers
    """

    if need_list[0] <= have_list[0]:
        return f'OK, строка: {string_number}'

    return f'Число больше чем нужно, строка: {string_number}'


def low(need_list, have_list):
    """
    Проверяет условие need_num_list[0] >= have_num_list[0]

    Принимает на входе числа, полученные функцией numbers
    """

    if need_list[0] >= have_list[0]:
        return f'OK, строка: {string_number}'

    return f'Число меньше чем нужно, строка: {string_number}'


def in_interval_equally(need_list, have_list):
    """
    Проверяет условие min(need_num_list) <= have_num_list[0] <= max(need_num_list)
    а так же наличие ошибок в указанных условиях

    Принимает на входе числа, полученные функцией numbers
    """

    if len(need_list) == 2:

        if min(need_list) == need_list[0] and max(need_list) == need_list[1]:
            if min(need_list) <= have_list[0] <= max(need_list):
                return f'OK, строка: {string_number}'

            return f'Число не подходит под условия "не менее чем и не более чем", ' \
                   f'строка: {string_number} '

        else:
            if need_list[1] <= have_list[0] <= need_list[0]:
                return f'OK, но возможна ошибка в условии, лучше проверить, строка: {string_number}'

            return f'Число не подходит под условия "не менее чем и не более чем", строка: {string_number}'

    elif len(need_list) == len(have_list) == 1:

        if need_list[0] == have_list[0]:
            return f'OK, но возможна ошибка в условии, лучше проверить, строка: {string_number}'

        return f'Число не подходит под условия "не менее чем и не более чем", строка: {string_number}'

    return f'Число не подходит под условия "не менее чем и не более чем", строка: {string_number}'


def in_interval_not_equally(need_list, have_list):
    """
    Проверяет условие min(need_num_list) < have_num_list[0] < max(need_num_list)
    а так же наличие ошибок в указанных условиях

    Принимает на входе числа, полученные функцией numbers
    """

    if min(need_list) == need_list[0] and max(need_list) == need_list[1]:
        if min(need_list) < have_list[0] < max(need_list):
            return f'OK, строка: {string_number}'

        return f'Число не подходит под условия "не менее чем и не более чем", строка: {string_number} '

    else:
        if need_list[1] < have_list[0] < need_list[0]:
            return f'OK, но возможна ошибка в условии, лучше проверить, строка: {string_number}'

        return f'Число не подходит под условия "не менее чем и не более чем", строка: {string_number}'


def in_interval_biggest(need_list, have_list):
    """
    Проверяет условие min(need_num_list) < have_num_list[0] <= max(need_num_list)
    а так же наличие ошибок в указанных условиях

    Принимает на входе числа, полученные функцией numbers
    """

    if min(need_list) == need_list[0] and max(need_list) == need_list[1]:
        if min(need_list) < have_list[0] <= max(need_list):
            return f'OK, строка: {string_number}'

        return f'Число не подходит под условия "< & <=", строка: {string_number}'

    else:
        if need_list[1] <= have_list[0] <= need_list[0]:
            return f'OK, но возможна ошибка в условии, лучше проверить, строка: {string_number}'

        return f'Число не подходит под условия "< & <=", строка: {string_number}'


def in_interval_lowest(need_list, have_list):
    """
    Проверяет условие min(need_num_list) <= have_num_list[0] < max(need_num_list)
    а так же наличие ошибок в указанных условиях

    Принимает на входе числа, полученные функцией numbers
    """

    if min(need_list) == need_list[0] and max(need_list) == need_list[1]:
        if min(need_list) <= have_list[0] < max(need_list):
            return f'OK, строка: {string_number}'

        return f'Число не подходит под условия "<= & <", строка: {string_number}'

    else:
        if need_list[1] <= have_list[0] <= need_list[0]:
            return f'OK, но возможна ошибка в условии, лучше проверить, строка: {string_number}'

        return f'Число не подходит под условия "<= & <", строка: {string_number}'


def lowest(need_list, have_list):
    """
    Проверяет условие need_num_list[0] > have_num_list[0]

    Принимает на входе числа, полученные функцией numbers
    """
    if need_list[0] > have_list[0]:
        return f'OK, строка: {string_number}'

    return f'Число больше чем нужно, строка: {string_number}'


def biggest(need_list, have_list):
    """
    Проверяет условие need_num_list[0] < have_num_list[0]

    Принимает на входе числа, полученные функцией numbers
    """

    if need_list[0] < have_list[0]:
        return f'OK, строка: {string_number}'

    return f'Число меньше чем нужно, строка: {string_number}'


def only_numbers(need_list, have_list):
    """
    Функция проверяет находятся ли в условии только числа (или интервал значений)

    Принимает на входе числа, полученные функцией numbers
    """
    if len(have_list) == 1 and len(need_list) == 2:
        if need_list[0] <= have_list[0] <= need_list[1]:
            return f'OK, строка: {string_number}'

        return f'Надо проверить, строка: {string_number}'

    return f'Надо проверить, строка: {string_number}'


def string_conditions(need_text, have_text):
    """
    Проверяет текстовое условие

    Принимает на входе числа, полученные функцией numbers
    """
    need_conditions_list = need_text.lower().replace(',', '').split()
    have_conditions_list = have_text.lower().replace(',', '').split()
    if 'или' in need_conditions_list and 'и' not in need_conditions_list:
        result = []
        for word1 in need_conditions_list:
            for word2 in have_conditions_list:
                if word1 == word2:
                    result.append(word1)
                    break

        if set(result) == set(have_conditions_list):
            return f'OK, строка: {string_number}'

        elif len((set(result)) - set(have_conditions_list)) >= 1:
            return f'Возможна опечатка в текстовом условии, строка: {string_number}'

        return f'Текстовое условие не выполнено, строка: {string_number}'

    elif 'и' in need_conditions_list:
        return f'Надо проверить, текстовое условие "И" строка: {string_number}'

    return f'Надо проверить, строка: {string_number}'


# интерфейс программы

layout = [
    [sg.Text('Файл заявки'), sg.InputText(), sg.FileBrowse('Указать путь к файлу')],
    [sg.Output(size=(88, 20))],
    [sg.Submit('Старт'), sg.Cancel('Отмена')]
]
window = sg.Window('СУПЕР ПРОВЕРЯЛЬЩИК ЗАЯВОК 3000', layout, background_color='green')
while True:
    event, values = window.read()
    # print(event, values) #debug
    if event in (None, 'Exit', 'Cancel', 'Отмена'):
        break
    if event == 'Submit' or event == 'Старт':
        # сама программа

        wordDoc = docx.Document(values[0])

        table_size = len(wordDoc.tables[1].rows)  # узнаем сколько во второй таблице в документе строк

        print('Начинаю проверку заявки. Строк к проверке:', table_size)

        for string in range(1, table_size):
            need = wordDoc.tables[1].rows[string].cells[2].text.lower().replace(' ', '')
            have = wordDoc.tables[1].rows[string].cells[3].text.lower().replace(' ', '')
            string_number = wordDoc.tables[1].rows[string].cells[0].text

            # выделяем числа из выбранных ячеек в отдельные списки при помощи функции numbers
            need_num_list = numbers(need)
            have_num_list = numbers(have)

            if (any(word in wordDoc.tables[1].rows[string].cells[1].text.lower()
                    for word in ('диап', 'диоп', 'интер'))):
                # интервал/диапазон в описании
                print(f'Указан какой-то интервал, необходимо согласовать с заказчиком, строка: {string_number}')

            if need == have:
                print(f'OK, строка: {string_number}')

            elif need == '' or have == '':
                print(f'Пустая ячейка, строка: {string_number}')

            elif (any(word in need for word in ('неменее', '≥', 'неменьше'))
                  and all(word not in need for word in ('неболее', '≤', 'небольше', '<'))):

                # не менее чем (need[0] <= have[0])
                print(more(need_num_list, have_num_list))

            elif (any(word in need for word in ('неболее', '≤', 'небольше'))
                  and all(word not in need for word in ('неменее', '≥', 'неменьше', '>'))):

                # не более чем (need[0] >= have[0])
                print(low(need_num_list, have_num_list))

            elif (any(word in need for word in ('неболее', '≤', 'небольше'))
                  and any(word in need for word in ('неменее', '≥', 'неменьше'))):

                # не менее чем и не более чем (need[0] <= have[0] <= need[1])
                print(in_interval_equally(need_num_list, have_num_list))

            elif (any(word in need for word in ('более', '>', 'больше'))
                  and any(word in need for word in ('неболее', '≤', 'небольше'))):

                # "< & <=" (need[0] < have[0] <= need[1])
                print(in_interval_biggest(need_num_list, have_num_list))

            elif (any(word in need for word in ('менее', '<', 'меньше'))
                  and any(word in need for word in ('неменее', '≥', 'неменьше'))):

                # "<= & <" (need[0] < have[0] <= need[1])
                print(in_interval_lowest(need_num_list, have_num_list))

            elif (any(word in need for word in ('менее', '<', 'меньше'))
                  and all(word not in need for word in ('более', '>', 'больше', '≥'))):

                # "менее чем" (need[0] > have[0])
                print(lowest(need_num_list, have_num_list))

            elif (any(word in need for word in ('более', '>', 'больше'))
                  and all(word not in need for word in ('менее', '<', 'меньше', '≤'))):

                # более чем (need[0] < have[0])
                print(biggest(need_num_list, have_num_list))

            elif (any(word in need for word in ('более', '>', 'больше'))
                  and any(word in need for word in ('менее', '<', 'меньше'))):

                # более чем и менее чем (need[0] < have[0] < need[1])
                print(in_interval_not_equally(need_num_list, have_num_list))

            elif 0 < len(need_num_list) <= 2:
                # в условии только числа
                print(only_numbers(need_num_list, have_num_list))

            elif len(set(list(need)) - set(list(have))) >= 1:
                # ошибки/опечатки и условия с и/или
                need = wordDoc.tables[1].rows[string].cells[2].text.lower()
                have = wordDoc.tables[1].rows[string].cells[3].text.lower()

                if ' и ' in need or ' или ' in need:
                    print(string_conditions(need, have))
                else:
                    print(f'Возможна ошибка или опечатка, строка: {string_number}')

            elif len(need) > len(have):
                # и/или
                print(string_conditions(need, have))

            else:
                print(f'Я не знаю что делать с этой строкой: {string_number}')

        print('Проверка файла завершена. Количество проверенных строк:', table_size)

    else:
        print('Что-то пошло не так')
